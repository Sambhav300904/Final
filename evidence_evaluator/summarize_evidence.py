"""
Summarize user-uploaded evidence files (PDF, DOCX, TXT, CSV, images).

Place files under evidences/<email>/ then run:
    python summarize_evidence.py --user engineer@company.com --days 7

Outputs JSON (+ combined .txt) under output/<email_safe>/.
"""

from __future__ import annotations

import argparse
import base64
import io
import json
import os
import re
import sys
import time
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

import httpx
import pandas as pd
from docx import Document
from dotenv import load_dotenv
from huggingface_hub import InferenceClient
from PIL import Image
from pypdf import PdfReader

MODULE_DIR = Path(__file__).resolve().parent
REPO_ROOT = MODULE_DIR.parent
ROOT_ENV = REPO_ROOT / ".env"
SYNC_ENV = REPO_ROOT / "cursor_mysql_sync" / ".env"

if ROOT_ENV.is_file():
    load_dotenv(ROOT_ENV)
if SYNC_ENV.is_file():
    load_dotenv(SYNC_ENV, override=False)

if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

EVIDENCES_DIR = MODULE_DIR / "evidences"
OUTPUT_DIR = MODULE_DIR / "output"
MAX_IMAGE_SIDE = 768
MAX_TEXT_CHARS = 12000

EMAIL_RE = re.compile(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$")

SUPPORTED_EXT = {
    ".pdf",
    ".docx",
    ".txt",
    ".csv",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".bmp",
}

IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}

PREVALIDATE_MAX_CHARS = int(os.getenv("EVIDENCE_PREVALIDATE_MAX_CHARS", "8000"))

RETRY_EXCEPTIONS = (
    httpx.ConnectTimeout,
    httpx.ReadTimeout,
    httpx.ConnectError,
    httpx.RemoteProtocolError,
    httpx.WriteTimeout,
)


def _optional(name: str, default: str) -> str:
    return os.getenv(name) or default


def resolve_period(days: int, start: date | None, end: date | None) -> tuple[date, date]:
    if end is None:
        end = date.today() - timedelta(days=1)
    if start is None:
        start = end - timedelta(days=max(days, 1) - 1)
    if start > end:
        raise ValueError("start must be <= end")
    return start, end


def _make_client() -> tuple[InferenceClient, str, str]:
    token = os.getenv("NOVITA_API_KEY") or os.getenv("HF_TOKEN")
    if not token:
        raise SystemExit(
            "Set HF_TOKEN or NOVITA_API_KEY in repo root .env for evidence summarization."
        )
    provider = _optional("EVIDENCE_LLM_PROVIDER", "novita")
    model = _optional("EVIDENCE_LLM_MODEL", "Qwen/Qwen3-VL-8B-Instruct")
    timeout = int(_optional("EVIDENCE_LLM_TIMEOUT_SEC", "600"))
    client = InferenceClient(token=token, provider=provider, timeout=timeout)
    return client, provider, model


def read_pdf(path: str) -> str:
    reader = PdfReader(path)
    chunks = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            text = f"[Error extracting page {i}: {exc}]"
        if text.strip():
            chunks.append(f"--- Page {i} ---\n{text.strip()}")
    if not chunks:
        return "[PDF contained no extractable text. It may be a scanned image PDF.]"
    return "\n\n".join(chunks)


def read_docx(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for t_idx, table in enumerate(doc.tables, start=1):
        parts.append(f"\n[Table {t_idx}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    if not parts:
        return "[Word document contained no readable text.]"
    return "\n".join(parts)


def read_txt(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def build_csv_summary(path: str) -> str:
    df = pd.read_csv(path)
    parts = [
        f"File: {os.path.basename(path)}",
        f"Rows: {len(df)}, Columns: {list(df.columns)}",
    ]
    if "Date" in df.columns:
        try:
            dates = pd.to_datetime(df["Date"], errors="coerce")
            parts.append(f"Date range: {dates.min()}  ->  {dates.max()}")
        except Exception:
            pass
    numeric_cols = df.select_dtypes(include="number").columns
    if len(numeric_cols):
        parts.append(
            "\nNumeric column statistics:\n" + df[numeric_cols].describe().to_string()
        )
    for col in df.select_dtypes(include="object").columns[:4]:
        parts.append(
            f"\nTop values for '{col}':\n" + df[col].value_counts().head(10).to_string()
        )
    parts.append("\nFirst 5 rows:\n" + df.head().to_string())
    parts.append("\nLast 5 rows:\n" + df.tail().to_string())
    return "\n".join(parts)


def encode_image(path: str, max_side: int = MAX_IMAGE_SIDE) -> str:
    img = Image.open(path).convert("RGB")
    w, h = img.size
    if max(w, h) > max_side:
        scale = max_side / max(w, h)
        img = img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=80, optimize=True)
    return "data:image/jpeg;base64," + base64.b64encode(buf.getvalue()).decode("utf-8")


def call_with_retries(fn, *args, label: str = "API"):
    last_exc = None
    for attempt in range(1, 4):
        try:
            return fn(*args)
        except RETRY_EXCEPTIONS as exc:
            last_exc = exc
            if attempt == 3:
                break
            wait = 5 * attempt
            print(
                f"{label} failed ({type(exc).__name__}), retrying in {wait}s ({attempt}/3)...",
                flush=True,
            )
            time.sleep(wait)
    raise last_exc  # type: ignore[misc]


def summarize_text_document(client, model: str, label: str, text: str) -> str:
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n[... truncated for length ...]"
    prompt = f"""You are a careful summarizer for R Systems AI proficiency evidence review.
The user uploaded a {label}.

Write a clear, well-structured summary. Include:
- Main topic and purpose (PR, ticket, training cert, screenshot, etc. if visible)
- Key facts, names, numbers, dates, links, or decisions
- Whether this looks like verifiable engineering evidence
- Short conclusion

Document content:
\"\"\"
{text}
\"\"\"
"""

    def _call():
        return client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": [{"type": "text", "text": prompt}]}],
            max_tokens=1500,
        )

    resp = call_with_retries(_call, label=f"{label} summary")
    return resp.choices[0].message.content or ""


def summarize_csv(client, model: str, path: str) -> str:
    structured = build_csv_summary(path)
    prompt = f"""You are given a structured summary of a CSV file (possible Cursor usage export).

{structured}

Write a clear natural-language summary:
- What the data represents
- Time range (if applicable)
- Notable patterns
- Short conclusion
"""

    def _call():
        return client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": [{"type": "text", "text": prompt}]}],
            max_tokens=1500,
        )

    resp = call_with_retries(_call, label="CSV summary")
    return resp.choices[0].message.content or ""


def summarize_image(client, model: str, path: str) -> str:
    image_messages = [{"type": "image_url", "image_url": {"url": encode_image(path)}}]
    prompt = (
        "Analyze this image as proficiency evidence. Describe what it shows and any "
        "visible text, links, PR numbers, tickets, or metrics. Note if it supports "
        "framework checklist evidence."
    )

    def _call():
        return client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": image_messages + [{"type": "text", "text": prompt}],
                }
            ],
            max_tokens=2000,
        )

    resp = call_with_retries(_call, label="Vision API")
    return resp.choices[0].message.content or ""


def extract_file_content(path: str) -> tuple[str, str]:
    """Local text extraction for pre-validation (no summarize LLM)."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return "pdf", read_pdf(path)
    if ext == ".docx":
        return "docx", read_docx(path)
    if ext == ".txt":
        return "txt", read_txt(path)
    if ext == ".csv":
        return "csv", build_csv_summary(path)
    if ext in IMAGE_EXT:
        return "image", ""
    if ext == ".doc":
        return "doc", ""
    return ext.lstrip(".") or "unknown", ""


def _trim_for_prevalidate(text: str) -> str:
    if len(text) <= PREVALIDATE_MAX_CHARS:
        return text
    return text[:PREVALIDATE_MAX_CHARS] + "\n\n[... truncated for pre-validation ...]"


def classify_image_prevalidate(client: InferenceClient, model: str, path: str) -> dict:
    """Vision-based relevance gate before full image summarize (source of truth for images)."""
    from validate_evidence import (
        _heuristic_validate,
        _parse_json_object,
        finalize_vision_prevalidate_verdict,
    )

    file_name = os.path.basename(path)
    prompt = (
        "You are validating uploaded IMAGE evidence for an AI Adoption proficiency pipeline.\n"
        "Look at the image and decide if it documents Cursor IDE, GitHub Copilot, or another "
        "AI coding-assistant / AI-assisted engineering workflow.\n\n"
        "IN SCOPE (is_relevant=true):\n"
        "- Cursor editor, Agent/Composer/Chat/Tabs, settings, rules, PRs with Cursor.\n"
        "- Cursor billing / usage pages: titles like \"Your usage\", cumulative spend charts, "
        "Group By Model — even when the word \"Cursor\" is not visible.\n"
        "- Model names that indicate Cursor billing dashboards: composer-1, composer-1.5, "
        "composer-2, composer-2.5, composer-2.5-fast, gpt-5.3-codex, gpt-5.3-codex-high, "
        "gpt-5.2-codex, claude-opus-4-7-thinking, claude-4.6-opus, claude-4.6-sonnet.\n"
        "- GitHub Copilot UI, usage, or training about Copilot.\n"
        "- Cursor analytics: work type / intent distribution / conversation insights charts.\n"
        "- Other AI coding tools (Codeium, Windsurf, Tabnine, Amazon Q Developer, etc.).\n\n"
        "OUT OF SCOPE (is_relevant=false):\n"
        "- Personal photos, unrelated utility/telecom bills, HR forms, payslips, timesheets.\n"
        "- Generic ChatGPT web UI, \"GPT Researcher\", localhost demos, or tools with no "
        "coding-assistant / Cursor / Copilot context.\n"
        "- Dashboards with zero AI-coding-tool or Cursor-model identifiers.\n\n"
        "Respond with ONLY a JSON object:\n"
        '{"is_relevant": bool, "confidence": number, '
        '"categories": ["cursor"|"copilot"|"other_ai_coding_tool"|"unrelated"|...], '
        '"reasoning": "<=2 sentences"}'
    )

    def _call():
        return client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": encode_image(path)}},
                        {"type": "text", "text": prompt},
                    ],
                }
            ],
            max_tokens=300,
        )

    try:
        resp = call_with_retries(_call, label="Image pre-validate")
        raw = resp.choices[0].message.content or ""
        parsed = _parse_json_object(raw)
        if isinstance(parsed, dict) and "is_relevant" in parsed:
            try:
                confidence = float(parsed.get("confidence", 0.0))
            except (TypeError, ValueError):
                confidence = 0.0
            confidence = max(0.0, min(1.0, confidence))
            categories = parsed.get("categories") or []
            if not isinstance(categories, list):
                categories = [str(categories)]
            categories = [str(c).strip().lower() for c in categories if str(c).strip()]
            is_relevant = bool(parsed.get("is_relevant"))
            min_conf = _validator_min_confidence()
            if is_relevant and confidence < min_conf:
                is_relevant = False
            verdict = {
                "is_relevant": is_relevant,
                "confidence": round(confidence, 3),
                "categories": categories or ["unknown"],
                "reasoning": str(parsed.get("reasoning") or "").strip(),
                "method": "llm_vision_prevalidate",
                "phase": "pre_validate",
            }
            return finalize_vision_prevalidate_verdict(
                verdict, file_name, categories=categories
            )
    except Exception as exc:
        heur = _heuristic_validate(file_name, file_name)
        heur["reasoning"] = (
            f"Vision pre-validate failed ({type(exc).__name__}: {exc}); "
            f"{heur['reasoning']}"
        )
        heur["method"] = "heuristic_fallback"
        heur["phase"] = "pre_validate"
        return heur

    heur = _heuristic_validate(file_name, file_name)
    heur["method"] = "heuristic_fallback"
    heur["phase"] = "pre_validate"
    return heur


def _validator_min_confidence() -> float:
    try:
        return float(os.getenv("EVIDENCE_VALIDATOR_MIN_CONFIDENCE") or "0.5")
    except ValueError:
        return 0.5


def _prevalidate_file(
    path: Path,
    *,
    validator_client,
    validator_model: str,
    summarize_client: InferenceClient,
    summarize_model: str,
) -> dict:
    """Run pre-summarize relevance check. Returns validation verdict dict."""
    from validate_evidence import is_enabled, validate_raw_extract

    name = path.name
    file_type, raw_text = extract_file_content(str(path))

    if not is_enabled():
        return {
            "is_relevant": True,
            "confidence": 1.0,
            "categories": ["unknown"],
            "reasoning": "Validator disabled; accepting file.",
            "method": "skipped",
            "phase": "pre_validate",
        }

    if file_type == "image":
        return classify_image_prevalidate(summarize_client, summarize_model, str(path))

    if file_type == "doc":
        return {
            "is_relevant": False,
            "confidence": 1.0,
            "categories": ["unknown"],
            "reasoning": "Old `.doc` format is not supported; re-save as `.docx`.",
            "method": "trivial",
            "phase": "pre_validate",
        }

    if not raw_text.strip():
        return {
            "is_relevant": False,
            "confidence": 1.0,
            "categories": ["unknown"],
            "reasoning": "No extractable text for pre-validation.",
            "method": "trivial",
            "phase": "pre_validate",
        }

    return validate_raw_extract(
        validator_client,
        validator_model,
        name,
        file_type,
        _trim_for_prevalidate(raw_text),
        min_confidence=_validator_min_confidence(),
    )


def summarize_one_file(client, model: str, path: str) -> tuple[str, str]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return "pdf", summarize_text_document(client, model, "PDF document", read_pdf(path))
    if ext == ".docx":
        return "docx", summarize_text_document(
            client, model, "Word (.docx) document", read_docx(path)
        )
    if ext == ".txt":
        return "txt", summarize_text_document(client, model, "text file", read_txt(path))
    if ext == ".csv":
        return "csv", summarize_csv(client, model, path)
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp"}:
        return "image", summarize_image(client, model, path)
    if ext == ".doc":
        return "doc", (
            "Old-style `.doc` files are not supported. Please re-save as `.docx`."
        )
    return ext or "unknown", f"Unsupported file type: `{ext}`."


def find_user_folder(email: str) -> tuple[Path | None, str | None]:
    email = email.strip().lower()
    if not email:
        return None, "Please provide --user with a valid email."
    if not EMAIL_RE.match(email):
        return None, f"`{email}` does not look like a valid email address."

    EVIDENCES_DIR.mkdir(parents=True, exist_ok=True)
    target = EVIDENCES_DIR / email
    if target.is_dir():
        return target, None

    available = sorted(
        d.name for d in EVIDENCES_DIR.iterdir() if d.is_dir()
    )
    hint = ""
    if available:
        hint = "\nAvailable folders:\n" + "\n".join(f"  - {d}" for d in available)
    else:
        hint = "\nThe evidences folder is empty."
    return None, f"No folder `evidences/{email}/`.{hint}"


def list_supported_files(folder: Path) -> list[Path]:
    files = []
    for path in sorted(folder.iterdir()):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXT:
            files.append(path)
    return files


def _rebuild_combined_from_summaries(
    email: str,
    folder: Path,
    period_start: date,
    period_end: date,
    entries: list[dict],
) -> str:
    parts = [
        f"Evidence summaries for {email}",
        f"Period: {period_start.isoformat()} to {period_end.isoformat()}",
        f"Source folder: {folder}",
        "",
    ]
    kept = False
    for entry in entries:
        if entry.get("error") or not (entry.get("summary") or "").strip():
            continue
        kept = True
        parts.append(f"### {entry.get('file')}")
        parts.append(entry.get("summary") or "")
        parts.append("")
    if not kept:
        parts.append(
            "(All uploaded evidence was rejected as not related to Cursor / "
            "Copilot / AI coding tool usage.)"
        )
    return "\n".join(parts).strip()


def _summarize_files(
    email: str,
    folder: Path,
    files: list[Path],
    period_start: date,
    period_end: date,
) -> dict | None:
    from validate_evidence import is_enabled, make_validator_client, validate_evidence_bundle

    client, provider, model = _make_client()
    validator_client, validator_provider, validator_model = make_validator_client()
    min_conf = _validator_min_confidence()

    print(f"Provider: {provider}, model: {model}", flush=True)
    if is_enabled():
        print(
            f"Validator: {validator_provider or 'heuristic'}, "
            f"model={validator_model or 'n/a'}, min_confidence={min_conf}",
            flush=True,
        )
    else:
        print("Validator: disabled (EVIDENCE_VALIDATOR_ENABLED)", flush=True)
    print(f"Files: {len(files)}", flush=True)

    accepted: list[dict] = []
    rejected: list[dict] = []
    files_rejected: list[str] = []
    files_failed: list[str] = []

    for i, path in enumerate(files, start=1):
        name = path.name
        print(f"[{i}/{len(files)}] {name} ...", flush=True)

        print("  pre-validate ...", flush=True)
        try:
            pre_verdict = _prevalidate_file(
                path,
                validator_client=validator_client,
                validator_model=validator_model,
                summarize_client=client,
                summarize_model=model,
            )
        except Exception as exc:
            pre_verdict = {
                "is_relevant": False,
                "confidence": 0.0,
                "categories": ["unknown"],
                "reasoning": f"Pre-validation error: {type(exc).__name__}: {exc}",
                "method": "error",
                "phase": "pre_validate",
            }

        if not pre_verdict.get("is_relevant"):
            tag = "REJECT"
            print(
                f"  -> {tag} pre-validate (conf={pre_verdict.get('confidence')}, "
                f"{pre_verdict.get('reasoning', '')[:80]})",
                flush=True,
            )
            rejected.append(
                {
                    "file": name,
                    "type": path.suffix.lower().lstrip(".") or "unknown",
                    "summary": "",
                    "error": None,
                    "validation": pre_verdict,
                }
            )
            files_rejected.append(name)
            continue

        print("  summarize ...", flush=True)
        t0 = time.perf_counter()
        try:
            file_type, summary = summarize_one_file(client, model, str(path))
            error = None
        except Exception as exc:
            file_type = path.suffix.lower().lstrip(".") or "unknown"
            summary = ""
            error = f"{type(exc).__name__}: {exc}"
        duration = round(time.perf_counter() - t0, 2)

        entry = {
            "file": name,
            "type": file_type,
            "duration_seconds": duration,
            "summary": summary,
            "error": error,
            "pre_validation": pre_verdict,
        }
        if error:
            files_failed.append(name)
            rejected.append({**entry, "validation": pre_verdict})
            print(f"  -> FAIL summarize: {error}", flush=True)
        else:
            accepted.append(entry)
            print(f"  -> summarized ({duration}s)", flush=True)

    bundle: dict = {
        "email": email,
        "period": {"start": period_start.isoformat(), "end": period_end.isoformat()},
        "evidence_folder": str(folder),
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "provider": provider,
        "model": model,
        "files_uploaded": [p.name for p in files],
        "files_processed": [e["file"] for e in accepted if not e.get("error")],
        "files_failed": files_failed,
        "files_rejected": files_rejected,
        "summaries": accepted,
        "rejected_summaries": rejected,
        "combined_evidence_text": "",
        "source": "user_upload",
    }

    if accepted and is_enabled():
        print(
            "Post-validating summaries (vision-approved images skip text re-check) ...",
            flush=True,
        )
        validate_evidence_bundle(
            bundle,
            client=validator_client,
            model=validator_model,
            min_confidence=min_conf,
        )

    bundle["combined_evidence_text"] = _rebuild_combined_from_summaries(
        email, folder, period_start, period_end, bundle.get("summaries") or []
    )

    n_kept = len(bundle.get("files_processed") or [])
    n_rej = len(bundle.get("files_rejected") or [])
    n_fail = len(bundle.get("files_failed") or [])
    print(
        f"Evidence done: {n_kept} kept, {n_rej} rejected, {n_fail} summarize failed "
        f"(of {len(files)} uploaded).",
        flush=True,
    )

    if not bundle.get("summaries"):
        return None
    return bundle


def load_evidence_bundle(
    email: str,
    period_start: date,
    period_end: date,
) -> dict | None:
    """Summarize relevant evidence if folder and files exist; otherwise return None."""
    folder, err = find_user_folder(email)
    if err or folder is None:
        return None
    files = list_supported_files(folder)
    if not files:
        return None
    return _summarize_files(email, folder, files, period_start, period_end)


def summarize_user(
    email: str,
    period_start: date,
    period_end: date,
) -> dict:
    folder, err = find_user_folder(email)
    if err:
        raise SystemExit(err)

    files = list_supported_files(folder)
    if not files:
        raise SystemExit(
            f"Folder found: {folder}\n"
            f"No supported files. Types: {', '.join(sorted(SUPPORTED_EXT))}"
        )

    bundle = _summarize_files(email, folder, files, period_start, period_end)
    if bundle is None:
        raise SystemExit(
            f"Folder: {folder}\n"
            f"{len(files)} file(s) uploaded but none passed relevance validation. "
            "Upload Cursor/Copilot/AI coding evidence or set EVIDENCE_VALIDATOR_ENABLED=false."
        )
    return bundle


def save_bundle(bundle: dict) -> Path:
    email = bundle["email"].replace("@", "_at_")
    period = bundle["period"]
    out_dir = OUTPUT_DIR / email
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = f"evidence_{period['start']}_{period['end']}"
    json_path = out_dir / f"{stem}.json"
    txt_path = out_dir / f"{stem}.txt"
    json_path.write_text(json.dumps(bundle, indent=2, ensure_ascii=False), encoding="utf-8")
    txt_path.write_text(bundle.get("combined_evidence_text", ""), encoding="utf-8")
    return json_path


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Summarize evidence files under evidences/<email>/"
    )
    parser.add_argument("--user", required=True, help="User email (folder name)")
    parser.add_argument("--days", type=int, default=30)
    parser.add_argument("--start", type=lambda s: date.fromisoformat(s))
    parser.add_argument("--end", type=lambda s: date.fromisoformat(s))
    args = parser.parse_args()

    email = args.user.strip().lower()
    start, end = resolve_period(args.days, args.start, args.end)
    bundle = summarize_user(email, start, end)
    path = save_bundle(bundle)
    print(f"Saved: {path}")
    print(f"Combined text: {path.with_suffix('.txt')}")
    ok = len(bundle.get("files_processed") or [])
    rej = len(bundle.get("files_rejected") or [])
    fail = len(bundle.get("files_failed") or [])
    print(f"Done: {ok} kept, {rej} rejected, {fail} summarize failed.")
    return 0 if fail == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
